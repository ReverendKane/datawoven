# text_input_tab.py
"""
Text Input Mode Tab - Direct text input and document file loading
"""
import logging
from pathlib import Path
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTextEdit, QPushButton,
    QSplitter, QGroupBox, QProgressBar, QTabWidget, QLineEdit,
    QFileDialog, QMessageBox
)
from PySide6.QtCore import Qt, QTimer

# Document processing imports
import docx  # python-docx for .docx files
from striprtf.striprtf import rtf_to_text  # striprtf for .rtf files
from odf import text as odf_text, teletype
from odf.opendocument import load as odf_load  # odfpy for .odt files
from bs4 import BeautifulSoup  # beautifulsoup4 for .html files
import chardet  # chardet for encoding detection

_LOGGER = logging.getLogger(__name__)


class TextInputTab(QWidget):
    """Text Input Mode - Direct text paste/type and document file loading"""

    def __init__(self, parent, shared_components, metadata_panel):
        super().__init__(parent)
        self.parent = parent
        self.shared_components = shared_components
        self.metadata_panel = metadata_panel
        self._text_timer = None
        self._loaded_file_path = None  # Track currently loaded file

        self.init_ui()

    def init_ui(self):
        """Initialize the text input mode interface"""
        layout = QHBoxLayout(self)
        splitter = QSplitter(Qt.Horizontal)
        layout.addWidget(splitter)

        # Left panel - Text mode controls
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)

        # Text input instructions
        info_group = QGroupBox("Text Input")
        info_layout = QVBoxLayout(info_group)
        from PySide6.QtWidgets import QLabel
        info_label = QLabel(
            "Load text from files or paste/type content directly for processing and summarization.")
        info_label.setWordWrap(True)
        info_layout.addWidget(info_label)
        left_layout.addWidget(info_group)

        # INPUT GROUP - File loading section
        input_group = QGroupBox("Input")
        input_layout = QVBoxLayout(input_group)

        # File path input and browse button
        file_input_layout = QHBoxLayout()
        self.input_file_path = QLineEdit()
        self.input_file_path.setPlaceholderText("Select a document file to load...")
        self.input_file_path.setReadOnly(True)
        file_input_layout.addWidget(self.input_file_path)

        input_browse_btn = QPushButton("Browse")
        input_browse_btn.setCursor(Qt.PointingHandCursor)
        input_browse_btn.clicked.connect(self.browse_input_file)
        file_input_layout.addWidget(input_browse_btn)

        input_layout.addLayout(file_input_layout)
        left_layout.addWidget(input_group)

        # Shared components - summarization only (no metadata)
        summary_group, self.ai_provider, self.summary_style, self.summary_length, self.auto_summarize, self.ai_status = self.shared_components.create_summarization_group()
        left_layout.addWidget(summary_group)

        # Output group
        output_group, self.output_folder_input, text_browse_btn, self.save_btn, self.copy_btn = self.shared_components.create_output_group()
        left_layout.addWidget(output_group)

        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        left_layout.addWidget(self.progress_bar)

        left_layout.addStretch()
        splitter.addWidget(left_panel)

        # Right panel - Text input and processing
        right_panel = self.create_text_input_panel()
        splitter.addWidget(right_panel)

        splitter.setSizes([400, 1000])

    def create_text_input_panel(self):
        """Create the text input and processing panel"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # Tab widget for different views
        self.tab_widget = QTabWidget()
        layout.addWidget(self.tab_widget)

        # Raw text input tab
        self.input_edit = QTextEdit()
        self.input_edit.setPlaceholderText("Paste or type your text content here, or load from a file...")
        self.input_edit.textChanged.connect(self.on_text_input_changed)
        self.tab_widget.addTab(self.input_edit, "Raw Text")

        # Summary tab
        self.summary_text_edit = QTextEdit()
        self.summary_text_edit.setPlaceholderText("Summarized content will appear here...")
        self.tab_widget.addTab(self.summary_text_edit, "Summary")

        # Final markdown tab
        self.markdown_edit = QTextEdit()
        self.markdown_edit.setPlaceholderText("Final markdown content for export...")
        self.tab_widget.addTab(self.markdown_edit, "Final Markdown")

        # Processing buttons
        button_layout = QHBoxLayout()

        self.summarize_btn = QPushButton("Summarize")
        self.summarize_btn.setCursor(Qt.PointingHandCursor)
        self.summarize_btn.setFixedHeight(40)
        self.summarize_btn.clicked.connect(lambda: self.parent.summarize_text('text'))
        self.summarize_btn.setEnabled(False)
        button_layout.addWidget(self.summarize_btn)

        self.generate_markdown_btn = QPushButton("Generate Markdown")
        self.generate_markdown_btn.setCursor(Qt.PointingHandCursor)
        self.generate_markdown_btn.setFixedHeight(40)
        self.generate_markdown_btn.clicked.connect(lambda: self.parent.generate_final_markdown('text'))
        self.generate_markdown_btn.setEnabled(False)
        button_layout.addWidget(self.generate_markdown_btn)

        self.save_btn = QPushButton("Save Content Bundle")
        self.save_btn.setCursor(Qt.PointingHandCursor)
        self.save_btn.clicked.connect(lambda: self.parent.save_markdown('text'))
        self.save_btn.setEnabled(False)
        self.save_btn.setFixedHeight(40)
        button_layout.addWidget(self.save_btn)

        layout.addLayout(button_layout)

        return panel

    def browse_input_file(self):
        """Open file dialog to select a document for loading"""
        file_filter = (
            "All Supported Files (*.txt *.docx *.doc *.rtf *.odt *.html *.htm);;"
            "Text Files (*.txt);;"
            "Word Documents (*.docx *.doc);;"
            "Rich Text Format (*.rtf);;"
            "OpenDocument Text (*.odt);;"
            "HTML Files (*.html *.htm);;"
            "All Files (*)"
        )

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Document File",
            "",
            file_filter
        )

        if file_path:
            self.load_document_file(file_path)

    def load_document_file(self, file_path: str):
        """Load and extract text from a document file"""
        try:
            path = Path(file_path)

            if not path.exists():
                QMessageBox.warning(self, "File Not Found", f"File does not exist:\n{file_path}")
                return

            # Update UI
            self.input_file_path.setText(file_path)
            self._loaded_file_path = file_path
            self.parent.statusBar().showMessage(f"Loading {path.name}...")

            # Extract text based on file extension
            suffix = path.suffix.lower()

            if suffix == '.txt':
                text = self._load_txt_file(path)
            elif suffix == '.docx':
                text = self._load_docx_file(path)
            elif suffix == '.doc':
                text = self._load_doc_file(path)
            elif suffix == '.rtf':
                text = self._load_rtf_file(path)
            elif suffix == '.odt':
                text = self._load_odt_file(path)
            elif suffix in ['.html', '.htm']:
                text = self._load_html_file(path)
            else:
                QMessageBox.warning(
                    self,
                    "Unsupported Format",
                    f"File format '{suffix}' is not supported.\n\n"
                    "Supported formats: .txt, .docx, .doc, .rtf, .odt, .html"
                )
                return

            if text:
                # Load text into the Raw Text pane
                self.input_edit.setPlainText(text)

                # Auto-populate metadata if available
                if hasattr(self.metadata_panel, 'title_input'):
                    if not self.metadata_panel.title_input.text().strip():
                        self.metadata_panel.title_input.setText(path.stem)

                if hasattr(self.metadata_panel, 'original_source_input'):
                    self.metadata_panel.original_source_input.setText(file_path)

                if hasattr(self.metadata_panel, 'source_type_combo'):
                    # Set source type based on file type
                    if suffix == '.html' or suffix == '.htm':
                        self.metadata_panel.source_type_combo.setCurrentText("html_file")
                    elif suffix in ['.docx', '.doc', '.rtf', '.odt']:
                        self.metadata_panel.source_type_combo.setCurrentText("document")
                    else:
                        self.metadata_panel.source_type_combo.setCurrentText("text_file")

                word_count = len(text.split())
                self.parent.statusBar().showMessage(
                    f"Loaded {path.name} - {word_count:,} words", 5000
                )

                _LOGGER.info(f"Successfully loaded file: {file_path} ({word_count} words)")
            else:
                QMessageBox.warning(
                    self,
                    "Empty File",
                    f"No text content could be extracted from:\n{path.name}"
                )

        except Exception as e:
            error_msg = f"Error loading file: {str(e)}"
            _LOGGER.error(error_msg, exc_info=True)
            QMessageBox.critical(self, "Load Error", error_msg)
            self.parent.statusBar().showMessage("File load failed", 5000)

    def _load_txt_file(self, path: Path) -> str:
        """Load plain text file with encoding detection"""
        try:
            # Try UTF-8 first
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Detect encoding
                with open(path, 'rb') as f:
                    raw_data = f.read()
                    result = chardet.detect(raw_data)
                    encoding = result['encoding'] or 'utf-8'

                _LOGGER.info(f"Detected encoding: {encoding}")
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()

        except Exception as e:
            _LOGGER.error(f"Error loading TXT file: {e}")
            raise

    def _load_docx_file(self, path: Path) -> str:
        """Load Microsoft Word .docx file"""
        try:
            doc = docx.Document(path)

            # Extract all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return '\n\n'.join(text_parts)

        except Exception as e:
            _LOGGER.error(f"Error loading DOCX file: {e}")
            raise

    def _load_doc_file(self, path: Path) -> str:
        """Load legacy Microsoft Word .doc file"""
        try:
            import win32com.client

            # Use COM automation on Windows
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                doc = word.Documents.Open(str(path.absolute()))
                text = doc.Content.Text
                doc.Close(False)
                return text
            finally:
                word.Quit()

        except ImportError:
            # Fallback: Convert to docx using LibreOffice if available
            try:
                import subprocess
                import tempfile

                with tempfile.TemporaryDirectory() as tmpdir:
                    # Convert to docx using LibreOffice
                    subprocess.run([
                        'soffice',
                        '--headless',
                        '--convert-to', 'docx',
                        '--outdir', tmpdir,
                        str(path)
                    ], check=True, capture_output=True)

                    docx_path = Path(tmpdir) / f"{path.stem}.docx"
                    return self._load_docx_file(docx_path)

            except (subprocess.SubprocessError, FileNotFoundError):
                raise Exception(
                    "Cannot load .doc files. Please install either:\n"
                    "- Microsoft Word (Windows)\n"
                    "- LibreOffice (for conversion to .docx)"
                )
        except Exception as e:
            _LOGGER.error(f"Error loading DOC file: {e}")
            raise

    def _load_rtf_file(self, path: Path) -> str:
        """Load Rich Text Format file"""
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()

            # Convert RTF to plain text
            text = rtf_to_text(rtf_content)
            return text

        except Exception as e:
            _LOGGER.error(f"Error loading RTF file: {e}")
            raise

    def _load_odt_file(self, path: Path) -> str:
        """Load OpenDocument Text file"""
        try:
            doc = odf_load(str(path))

            # Extract all text elements
            text_parts = []
            all_paragraphs = doc.getElementsByType(odf_text.P)

            for paragraph in all_paragraphs:
                text_content = teletype.extractText(paragraph)
                if text_content.strip():
                    text_parts.append(text_content)

            return '\n\n'.join(text_parts)

        except Exception as e:
            _LOGGER.error(f"Error loading ODT file: {e}")
            raise

    def _load_html_file(self, path: Path) -> str:
        """Load HTML file and extract text content"""
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()

            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text

        except Exception as e:
            _LOGGER.error(f"Error loading HTML file: {e}")
            raise

    def on_text_input_changed(self):
        """Handle text input changes"""
        has_text = bool(self.input_edit.toPlainText().strip())
        self.summarize_btn.setEnabled(has_text)

        # Only auto-summarize if enabled, has content, and not already processing
        if (has_text and
            self.auto_summarize.isChecked() and
            (not hasattr(self.parent, 'summarization_thread') or
             self.parent.summarization_thread is None or
             not self.parent.summarization_thread.isRunning())):

            # Cancel any existing timer
            if self._text_timer is not None:
                self._text_timer.stop()

            # Create new timer
            self._text_timer = QTimer()
            self._text_timer.setSingleShot(True)
            self._text_timer.timeout.connect(lambda: self.parent.summarize_text('text'))
            self._text_timer.start(1000)

    def reset_fields(self):
        """Reset all fields and content"""
        # Clear text areas
        self.input_edit.clear()
        self.summary_text_edit.clear()
        self.markdown_edit.clear()

        # Clear file path
        self.input_file_path.clear()
        self._loaded_file_path = None

        # Reset buttons
        self.summarize_btn.setEnabled(False)
        self.generate_markdown_btn.setEnabled(False)
        self.save_btn.setEnabled(False)

        # Switch back to first tab
        self.tab_widget.setCurrentIndex(0)

        self.parent.statusBar().showMessage("Text mode reset - ready for new content")
